import os
from docxtpl import DocxTemplate, RichText, Subdoc
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from typing import Any, Dict, List, Union, Optional

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESULT_DIR = os.path.join(BASE_DIR, "result")

ALIGN_MAP = {
  "left": WD_ALIGN_PARAGRAPH.LEFT,
  "center": WD_ALIGN_PARAGRAPH.CENTER,
  "right": WD_ALIGN_PARAGRAPH.RIGHT,
  "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

def set_run_style(run, bold=False, italic=False, underline=False, font_size: Optional[int]=None, font_name: str="Times New Roman"):
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if font_size:
        run.font.size = Pt(font_size)
    run.font.name = font_name

def indent(p: Paragraph, left_in: float = 0.25):
    fmt = p.paragraph_format
    fmt.left_indent = Inches(left_in)

def add_para(doc: Document, text: List[Dict], font_size=12, align=None, space_before=None, space_after=None, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for item in text:
        t = item.get("text", "")
        bold = item.get("bold", False)
        italic = item.get("italic", False)
        underline = item.get("underline", False)
        r = p.add_run(t)
        set_run_style(r, bold=bold, italic=italic, underline=underline, font_size=font_size)
    if align:
        p.alignment = ALIGN_MAP[align]
    if space_before:
        fmt = p.paragraph_format
        fmt.space_before = Pt(float(space_before))
    if space_after:
        fmt = p.paragraph_format
        fmt.space_after = Pt(float(space_after))
    return p

def add_items(doc: Document, items: List[List[Dict]], bullet: Optional[str]=None):
    for i in items:
        p = doc.add_paragraph()
        if bullet:
            r = p.add_run(f"{bullet} ")
        for text in i:
            t = text.get("text", "")
            bold = text.get("bold", False)
            italic = text.get("italic", False)
            underline = text.get("underline", False)
            r = p.add_run(t)
            set_run_style(r, bold=bold, italic=italic, underline=underline, font_size=12)

        fmt = p.paragraph_format
        fmt.space_before = Pt(0)

def render_cell(cell, value: List[List[Dict]], font_size: Optional[int]=None):
    cell.text = ""
    first = True
    for i, line in enumerate(value):
        p = cell.paragraphs[0] if first and cell.paragraphs else cell.add_paragraph()
        first = False
        p.alignment = ALIGN_MAP.get(line.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        space_before = line.get("space_before", None)
        space_after = line.get("space_before", None)
        if space_before:
            fmt = p.paragraph_format
            fmt.space_before = Pt(float(space_before))

        if space_after:
            fmt = p.paragraph_format
            fmt.space_after = Pt(float(space_after))
        elif i == len(value) - 1:
            fmt = p.paragraph_format
            fmt.space_after = Pt(4)

        text = line.get("text", [])
        for item in text:
            t = item.get("text", "")
            bold = item.get("bold", False)
            italic = item.get("italic", False)        
            underline = item.get("underline", False)
            r = p.add_run(t)
            set_run_style(r, bold=bold, italic=italic, underline=underline, font_size=font_size)

def set_column_widths(table, widths):
    for col_idx, w in enumerate(widths):
        for r in table.rows:
            r.cells[col_idx].width = Inches(w)

def add_table(doc: Document, headers: List[Dict], row: List[List[Union[str, int, float]]], font_size: Optional[int]=None, col_widths: Optional[List[float]]=None):
    if not headers:
        return "Cần có headers hoặc ít nhất một hàng dữ liệu."
    
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.autofit = True
    table.style = "Table Grid"

    for col, h in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = p.paragraph_format
        fmt.space_after = Pt(4)
        t = h.get("text", "")
        bold = h.get("bold", False)
        italic = h.get("italic", False)
        underline = h.get("underline", False)
        r = p.add_run(t)
        set_run_style(r, bold=bold, italic=italic, underline=underline, font_size=font_size)

    for row_data in row:
        row_cells = table.add_row().cells
        for col_idx, value in enumerate(row_data):
            render_cell(row_cells[col_idx], value, font_size=font_size)

    if col_widths and len(col_widths) == ncols and sum(col_widths) <= 6.5:
        table.autofit = False
        set_column_widths(table, col_widths)

    return table

def init_doc(font_name: str="Times New Roman", font_size: int=12):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = font_name
    normal.font.size = Pt(font_size)

    para_fmt = normal.paragraph_format
    para_fmt.space_after = Pt(0)
    para_fmt.space_before = Pt(4)
    para_fmt.line_spacing = 1

    return doc

def save_doc(doc: Document, output_path: str):
    doc.save(output_path)
    return output_path

def gen_doc(plan: dict, output="giao_an.docx"):
    default_font_name = (plan.get("meta").get("font_name") or "Times New Roman")
    default_font_size = int(plan.get("meta").get("font_size") or 12)

    doc = init_doc(default_font_name, default_font_size)

    steps = plan.get("step", [])

    if steps:
        for func in steps:
            func_name = func.get("tool", "")
            if func_name:
                match func_name:
                    case "add_para":
                        args = func.get("args", "")
                        text = args.get("text", [])
                        font_size = args.get("font_size", default_font_size) 
                        align = args.get("align", "left")
                        space_before = args.get("space_before", None)
                        space_after = args.get("space_after", None)
                        add_para(doc, text, font_size, align, space_before, space_after)
                    case "add_items":
                        args = func.get("args", "")
                        items = args.get("items", [])
                        bullet = args.get("bullet", None)
                        add_items(doc, items, bullet)
                    case "add_table":
                        args = func.get("args", "")
                        headers = args.get("headers", [])
                        rows = args.get("rows", [])
                        font_size = args.get("font_size", default_font_size)
                        col_widths = args.get("col_widths", None)
                        add_table(doc, headers, rows, font_size, col_widths)
    
    output_path = os.path.join(RESULT_DIR, output)
    save_doc(doc, output_path)

    return output_path