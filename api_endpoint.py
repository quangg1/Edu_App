import os, json, threading, tempfile, time, io, re, logging, traceback, mimetypes, asyncio, uuid
import uvicorn
from fastapi import FastAPI, UploadFile, Form, Request, HTTPException,Body,File, Depends
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from fastapi.middleware import cors
from fastapi.middleware.cors import CORSMiddleware

from pydantic import BaseModel, HttpUrl,ConfigDict, Field, ValidationError
from pydantic.alias_generators import to_camel
from typing import Optional, List, Dict, Any, AsyncGenerator
from pathlib import Path

from google import genai
from google.genai.types import Tool, GenerateContentConfig, Part, ToolConfig, FunctionCallingConfig, AutomaticFunctionCallingConfig
from google.genai import types
from docx import Document
import pdfplumber

from chatbot_mcq import generate_mcqs, extract_text_from_docx, extract_text_from_pdf
from modelMCQs import  extract_text_from_pdf, extract_text_from_docx, generate_test, export_quiz_to_docx, QuizzContentValidate
from Quiz_demo.modelv2 import extract_text_from_pdf, extract_text_from_docx, build_prompt, take_json, cut_question_in_buffer, export_json_to_md, convert_md_to_docx
from Quiz_demo.config import  QuizzContentValidate, setup_logging, QuizResponse, QuizPayLoad
from rubric.services.streaming import stream_rubric_gemini_direct, DOWNLOAD_STORE_RUBRIC
from rubric.services.prompt_builder import RUBRIC_SYSTEM_PROMPT_GEMINI, _format_user_block

from lessonplan.model.k12_model import K12
from lessonplan.model.kindergarten_model import Kindergarten
from lessonplan.model.create_doc_model import DocCreate
from lessonplan.doc_utils import gen_doc
from lessonplan.utils import convert_k12, convert_kindergarten, validate_k12, validate_kindergarten, has_valid_key, parse_json, has_valid_tool, parse_json_tool
from lessonplan.html_utils import convert_k12_objectives, convert_k12_resources, convert_k12_activities, convert_kindergarten_objectives, convert_kindergarten_resources, convert_kindergarten_activities, gen_html
from lessonplan.service import get_mime_type, generate_lessonplan_template, generate_lessonplan_custom

from dotenv import load_dotenv

load_dotenv()
gemini_api_key = os.getenv("GEMINI_API_KEY")
client = genai.Client(api_key=gemini_api_key)

app = FastAPI()
origins = [
    "https://edu-app-h2f9.onrender.com",  # Origin của frontend local
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"], 
    allow_headers=["*"],  
)
DOWNLOAD_STORE: Dict[str, str] = {}

# ----------------------------
# ✅ FastAPI Endpoint for Generating MCQs
# ----------------------------
@app.post("/generate_mcqs")
async def generate_mcqs_endpoint(
    file: UploadFile = Form(...),
    num_questions: int = Form(...),
    subject: str = Form(None),
    grade: str = Form(None),
    topic: str = Form(None)
):
    # 🕐 Bắt đầu đo thời gian xử lý
    start_time = time.time()

    with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as temp_file:
        temp_path = temp_file.name
        content = await file.read()
        temp_file.write(content)

    try:
        # 🧾 Đọc nội dung file
        if file.filename.endswith(".pdf"):
            text = extract_text_from_pdf(temp_path)
        elif file.filename.endswith(".docx"):
            text = extract_text_from_docx(temp_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF or DOCX file.")

        if not text.strip():
            raise HTTPException(status_code=400, detail="No text extracted from document.")

        # 🤖 Gọi model sinh câu hỏi
        mcq_result = generate_mcqs(text, num_questions=num_questions, subject=subject, grade=grade, topic=topic)


        end_time = time.time()
        response_payload = {
            "success": True,
            "provider": "Google Gemini",
            "model": "gemini-2.5-flash",
            "processing_time": round(end_time - start_time, 2),
            "subject": subject,
            "grade": grade,
            "topic": topic,
            "num_questions": num_questions,
            "questions": mcq_result.get("questions", []),
            "study_suggestions": mcq_result.get("study_suggestions", []),
            "references": mcq_result.get("references", []),
            "search_terms": mcq_result.get("search_terms", []),
        }

        return JSONResponse(content=response_payload, status_code=200)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating MCQs: {str(e)}")

    finally:
        os.remove(temp_path)

# ----------------------------
# ✅ FastAPI Endpoint for Generating quiz.docx
# ----------------------------
@app.post("/generate-quiz")
async def api_generate_quiz(
    file: Optional[UploadFile] = File(default=None),
    text_content: Optional[str] = Form(default=None),
    name: Optional[str] = Form(default=None),
    subject: Optional[str] = Form(default=None),
    grade: Optional[int] = Form(default=None),
    num_questions: int = Form(default=10),
    time_limit: int = Form(default=45),
    difficulty: str = Form(default="Medium"),
    topic: Optional[str] = Form(default=None),
    percentage: int = Form(default=70),
):
    start_time = time.time()

    quiz_data = QuizzContentValidate(
        text_content=text_content,
        name=name,
        subject=subject,
        grade=grade,
        num_questions=num_questions,
        time_limit=time_limit,
        difficulty=difficulty,
        topic=topic,
        percentage=percentage,
    )

    # Check input
    if not file and not (text_content and text_content.strip()):
        raise HTTPException(status_code=400, detail="Cần upload file hoặc cung cấp text_content.")
    
    # Extract text
    if file:
        suffix = Path(file.filename).suffix.lower()
        tmp = Path("uploads") / f"{int(time.time())}_{file.filename}"
        tmp.parent.mkdir(parents=True, exist_ok=True)
        with open(tmp, "wb") as f:
            f.write(await file.read())
        if suffix == ".pdf":
            text = extract_text_from_pdf(str(tmp))
        elif suffix == ".docx":
            text = extract_text_from_docx(str(tmp))
        else:
            tmp.unlink(missing_ok=True)
            raise HTTPException(status_code=415, detail="Chỉ hỗ trợ .pdf hoặc .docx")
        tmp.unlink(missing_ok=True)
    else:
        text = quiz_data.text_content.strip()

    # Sinh đề
    response = generate_test(
        text, 
        quiz_data.name, 
        quiz_data.subject, 
        quiz_data.grade, 
        quiz_data.num_questions, 
        quiz_data.time_limit, 
        quiz_data.difficulty, 
        quiz_data.topic, 
        quiz_data.percentage
    )

    if "error" in response:
        return JSONResponse(status_code=502, content=response)
    
    # Xuất ra file DOCX
    result_filename = export_quiz_to_docx(data=response) # 💡 Đổi tên biến từ 'result' thành 'result_filename' để rõ ràng hơn
    
    if not result_filename: # Sử dụng tên biến mới
        raise HTTPException(status_code=500, detail="Lỗi tạo DOCX!")

    duration = time.time() - start_time
    return {
        "status": "success",
        "name": quiz_data.name,
        "subject": quiz_data.subject,
        "grade": quiz_data.grade,
        "difficulty": quiz_data.difficulty,
        "num_questions": quiz_data.num_questions,
        "time_limit": quiz_data.time_limit,
        "duration": f"{duration:.2f}s",
        "file_name": result_filename,
        "questions": response["questions"],  # nếu hàm generate_test có trả danh sách câu hỏi
        "download_url": f"/download/{result_filename}"
    }
logger = setup_logging()
def sse(event: str, data: Any) -> str:
    """
    Format a Server-Sent Event (SSE) message.
    Each SSE event must contain "event:" and "data:" lines, then an empty line as a separator.
    The payload must be a string; we JSON-encode dict/list to ensure consistent client parsing.
    Args:
        event (str): The event type/name.
        data (Any): The event payload, can be str, dict, list, etc.
    Returns:
        str: Formatted SSE message.
    """
    payload = data if isinstance(data, str) else json.dumps(data, ensure_ascii=False)
    return f"event: {event}\n" f"data: {payload}\n\n"

@app.get("/healthz")
def healthz(): return "ok"

@app.post("/generate-quiz-stream")
async def api_generate_quiz(
    file: Optional[UploadFile] = File(default=None),
    text_content: Optional[str] = Form(default=None),
    name: Optional[str] = Form(default=None),
    subject: Optional[str] = Form(default=None),
    grade: Optional[int] = Form(default=None),
    num_questions: int = Form(default=5),
    time_limit: int = Form(default=15),
    difficulty: str = Form(default="Easy"),
    topic: Optional[str] = Form(default=None),
    percentage: int = Form(default=70),
):
    start_time= time.time()

    quiz_data = QuizzContentValidate(
        text_content=text_content,
        name=name,
        subject=subject,
        grade=grade,
        num_questions=num_questions,
        time_limit=time_limit,
        difficulty=difficulty,
        topic=topic,
        percentage=percentage,
    )
    
    # Check exist file or text:
    if not file and not (text_content and text_content.strip()):
        raise HTTPException(status_code=400, detail="Cần upload file hoặc cung cấp text_content.")
    
    # Extract input content
    if file:
        suffix = Path(file.filename).suffix.lower()
        tmp = Path("uploads")/f"{int(time.time())}_{file.filename}"
        tmp.parent.mkdir(parents=True, exist_ok=True)
        with open(tmp, "wb") as f:
            f.write(await file.read())
        if suffix == ".pdf":
            text = extract_text_from_pdf(str(tmp))
        elif suffix == ".docx":
            text = extract_text_from_docx(str(tmp))
        else:
            try:
                tmp.unlink(missing_ok=True)
            finally:
                pass
            raise HTTPException(status_code=415, detail="Chỉ hỗ trợ .pdf hoặc .docx")
        try: 
            tmp.unlink(missing_ok=True)
        finally:
            pass
    else:
        text = QuizzContentValidate.text_content.strip()

    prompt = build_prompt(
        text, 
        quiz_data.name, 
        quiz_data.subject, 
        quiz_data.grade, 
        quiz_data.num_questions, 
        quiz_data.time_limit, 
        quiz_data.difficulty, 
        quiz_data.topic, 
        quiz_data.percentage
    )
    
    async def gen() -> AsyncGenerator[str, None]:
        yield sse("status", {"message":"calling gemini....."})
        resp = client.models.generate_content_stream(
            model = "gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=("""
                - Trả về đúng json schema đã được định nghĩa
                - Sắp xếp câu hỏi theo thứ tự: Trắc nghiệm -> tự luận, Dễ -> khó
                - Response chỉ có json hợp lệ, KHÔNG in text ra ngoài JSON. 
                """),
                response_mime_type='application/json',
                response_json_schema=QuizResponse.model_json_schema()
            )
        )
        full_json = []
        buffer = ""
        sent_question_ids = set()
        for chunk in resp:
            t = getattr(chunk, "text", None)
            if not t:
                continue
            full_json.append(t)
            buffer += t
            
            
            new_question, buffer = cut_question_in_buffer(buffer)
            for q in new_question:
                q_id = q.get("id")
                if q_id in sent_question_ids:
                    continue
                sent_question_ids.add(q_id)
                yield sse("generating",{"detail": q})

        yield sse("status", {"message":"generated"})

        raw_text = "".join(full_json).strip()
        
        cleaned_text = raw_text
        

        if cleaned_text.startswith("```json"):
            cleaned_text = cleaned_text[len("```json"):]
        

        if cleaned_text.endswith("```"):
            cleaned_text = cleaned_text[:-len("```")]


        cleaned_text = cleaned_text.strip()
        
        quiz = take_json(cleaned_text)


        if isinstance(quiz, dict) and "error" in quiz:
            quiz["raw_output"] = raw_text
            logger.error(f"JSON Parsing Error. Raw output: {raw_text[:500]}...")

        yield sse("Done", {"detail": quiz})

    duration = time.time() - start_time
    logger.info(f"Time:{duration}s")
    return StreamingResponse(
        gen(), 
        media_type="text/event-stream",
        headers={
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "Transfer-Encoding": "chunked",
        "X-Accel-Buffering": "no",
        }
    )

#-----------------------------
# API Export DOCX FILE
#-----------------------------
@app.post('/export-docx')
async def Export_file_Docx(payload: QuizPayLoad):
    data = payload.quiz

    if "questions" not in data or not isinstance(data["questions"], list):
        raise HTTPException(status_code=400, detail="Quiz không hợp lệ: thiếu danh sách câu hỏi.")

    raw_md = export_json_to_md(data, "exports")
    if not raw_md:
        raise HTTPException(status_code=500, detail="Không thể tạo file MarkDown")
    
    docx_path = convert_md_to_docx(raw_md, "exports", "templates/template.docx")
    if not docx_path:
        raise HTTPException(status_code=500, detail="Không thể chuyển sang DOCX.")

    filename = Path(docx_path).name
    download_url = f"/api/v2/download/{filename}"

    return {
        "status": "ok",
        "download_url": download_url,
        "filename": filename,
    }


@app.post('/api/v2/download/{filename}')
async def download_doc(file_name: str):
    file_path = Path("exports") / file_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File không tồn tại")

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name
    )
@app.post(
    "/rubric/generate_gemini_stream",
    summary="Sinh Rubric JSON bằng Gemini (streaming SSE)",
    description="Upload 0..n file (PDF/DOCX/PPTX/TXT) -> đính vào Gemini -> stream tiến trình sinh JSON rubric.",
)
async def generate_rubric_gemini_stream(
    rubric_title: str = Form(..., description="Tiêu đề Rubric", examples=["Bài thuyết trình tính chất hóa học của kim loại"]),
    subject: str = Form(..., description="Môn học", examples=["Khoa học tự nhiên (Môn Hóa)"]),
    grade_level: str = Form(..., description="Khối lớp", examples=["Lớp 9"]),
    assessment_type: str = Form(..., description="Loại đánh giá", examples=["Thuyết trình"]),
    number_of_criteria: int = Form(..., description="Số tiêu chí", examples=[5]),
    user_prompt: Optional[str] = Form("", description="Yêu cầu bổ sung (nếu có)"),
    files: Optional[List[UploadFile]] = File(None, description="Tài liệu tham khảo (tuỳ chọn)"),
):
    """
    Streaming endpoint for generating rubric JSON with Gemini.
    Emits real-time progress and partial outputs as SSE events.
    Args:
        rubric_title (str): Title of the rubric.
        subject (str): Subject name.
        grade_level (str): Grade level.
        assessment_type (str): Type of assessment.
        number_of_criteria (int): Number of criteria.
        user_prompt (str): Additional user prompt.
        files (Optional[List[UploadFile]]): List of user-uploaded files.
        model (str): Gemini model to use.
    Yields:
        AsyncGenerator[str, None]: SSE events as strings.
    """
    gen = stream_rubric_gemini_direct(
        rubric_title=rubric_title,
        subject=subject,
        grade_level=grade_level,
        assessment_type=assessment_type,
        number_of_criteria=number_of_criteria,
        user_prompt=user_prompt,
        files=files,
    )
    return StreamingResponse(gen
                             , media_type="text/event-stream"
                             , headers={
        "Cache-Control": "no-cache"
        , "Connection": "keep-alive"
        , "X-Accel-Buffering": "no"
                             })

@app.get("/download-rubric/{token}")
async def download_rubric(token: str):
    """
    Download generated lesson plan by token.
    Args:
        token (str): Unique token for the rubric.
    Returns:
        FileResponse: The rubric file response.
    """
    meta = DOWNLOAD_STORE.get(token)
    if not meta or not os.path.exists(meta["path"]):
        raise HTTPException(status_code=404, detail="Invalid or expired token")

    return FileResponse(
        meta["path"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=meta["filename"],
    )
# ----------------------------
# ✅ FastAPI Endpoint for Generating Lesson Plan
# ----------------------------
@app.post("/lessonplan/k12")
async def generate_k12(
    prompt: str = Form(..., description="Instruction or prompt for the generator"),
    files: Optional[List[UploadFile]] = File(default=None, description="Optional files to include")
):
    try:
        SYS_MSG = """
            You are an expert lesson plan generator. Your task is to generate structured lesson plans based on user requirements, strictly following the specified JSON schema with markdown content.

            # User Instruction: {instruction}

            # Tone: formal, clear, structured

            - All content must be in plain text. Do not use any markdown formatting such as bold, italic, headings, lists, or other special symbols.
            - All text content must be in Vietnamese.
            - All lesson plans must follow the correct logical flow for a typical teaching period in Vietnam.
            - For each activity, clearly describe actions of both Teacher (Giáo viên) and Student (Học sinh).
            - Focus on clear, measurable objectives and expected learning outcomes.
            - Make sure that the lesson plan is pedagogically sound, well-structured, and matches the educational context of Vietnam's K-12 education system.
            - All generated Vietnamese text must use correct spelling, grammar, and punctuation.
            - If quotation marks are used, always format them using Word's typographic quotation marks (“ ”)
            - Do not include images or external URLs.
            - User instrction should always be followed and should supercede any other instruction.
            - Output must be a valid JSON object strictly following the LessonPlan schema.
        """
        response_format = K12.model_json_schema()
        system_prompt = SYS_MSG.format(instruction=prompt)
        contents = [prompt]
        for f in files:
            data = await f.read()
            mime = f.content_type or "application/octet-stream"
            contents.append(Part.from_bytes(
                data=data,
                mime_type=mime,
            ))
        
        error_msg = None
        for attempt in range(3):
            response = generate_lessonplan_template(contents, system_prompt, response_format)
            response_json = json.loads(response)
            try:
                validate = validate_k12(response_json)
                with open("lessonplan/temp/out_k12.txt", "w", encoding="utf-8") as f:
                    f.write(response)
                output_path = convert_k12(response_json)
                if output_path:
                    return FileResponse(
                        path=output_path,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename="giao_an.docx"
                    )
                else:
                    return JSONResponse(
                        status_code=500,
                        content={"error": "Đã xảy ra lỗi khi tạo file docx."}
                    )
            except ValidationError as e:
                error_msg = str(e.errors())
                print(f"Attempt #{attempt+1}: Error")

        if error_msg:
            return JSONResponse(
                status_code=500,
                content={"error": "Đã xảy ra lỗi khi tạo giáo án.", "details": str(error_msg)}
            )

    except Exception as e:
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": "Đã xảy ra lỗi trong quá trình xử lý.", "details": str(e)}
        )

@app.post("/lessonplan/kindergarten")
async def generate_kindergarten(
    prompt: str = Form(..., description="Instruction or prompt for the generator"),
    files: Optional[List[UploadFile]] = File(default=None, description="Optional files to include")
):
    try:
        SYS_MSG = """
            You are an expert kindergarten lesson plan generator specialized in Vietnam's STEM-5E education framework. Your task is to generate structured kindergarten lesson plans based on user requirements, strictly following the specified JSON schema with markdown content.

            # User Instruction: {instruction}

            # Tone: formal, clear, structured

            - All text must be written in plain text without any Markdown formatting such as bold, italics, headings, lists, or special symbols.
            - All text content must be written in Vietnamese using correct spelling, grammar, and punctuation.
            - All lesson plans must follow the correct logical flow for a typical teaching period in Vietnam.
            - For each activity, clearly describe actions of both Teacher (Giáo viên) and Student (Học sinh).
            - Focus on clear, measurable objectives and expected learning outcomes.
            - Make sure that the lesson plan is pedagogically sound, well-structured, and matches the educational context of Vietnam's K-12 education system.
            - All generated Vietnamese text must use correct spelling, grammar, and punctuation.
            - If quotation marks are used, always format them using Word's typographic quotation marks (“ ”)
            - Do not include images or external URLs.
            - User instrction should always be followed and should supercede any other instruction.
            - Output must be a valid JSON object strictly following the LessonPlan schema.
        """
        response_format = Kindergarten.model_json_schema()
        system_prompt = SYS_MSG.format(instruction=prompt)
        contents = [prompt]
        for f in files:
            data = await f.read()
            mime = f.content_type or "application/octet-stream"
            contents.append(Part.from_bytes(
                data=data,
                mime_type=mime,
            ))

        error_msg = None
        for attempt in range(3):
            response = generate_lessonplan_template(contents, system_prompt, response_format)
            response_json = json.loads(response)
            try:
                validate = validate_kindergarten(response_json)
                with open("lessonplan/temp/out_kindergarten.txt", "w", encoding="utf-8") as f:
                    f.write(response)
                output_path = convert_kindergarten(response_json)
                if output_path:
                    return FileResponse(
                        path=output_path,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename="giao_an.docx"
                    )
                else:
                    return JSONResponse(
                        status_code=500,
                        content={"error": "Đã xảy ra lỗi khi tạo file docx."}
                    )
            except ValidationError as e:
                error_msg = str(e.errors())
                print(f"Attempt #{attempt+1}: Error")

        if error_msg:
            return JSONResponse(
                status_code=500,
                content={"error": "Đã xảy ra lỗi khi tạo giáo án.", "details": str(error_msg)}
            )

    except Exception as e:
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": "Đã xảy ra lỗi trong quá trình xử lý.", "details": str(e)}
        )

@app.post("/lessonplan/custom")
async def generate_custom(
    prompt: str = Form(..., description="Instruction or prompt for the generator"),
    files: List[UploadFile] = File(default=None, description="Optional files to include")
):
    try:
        SYS_MSG = """
            You are an expert Word content planner specialized in structured document generation. 
            Your task is to output a valid JSON object strictly following the DocCreate schema based on the user instruction below.

            # User Instruction: {instruction}

            # Tone: formal, clear, structured

            - You must only return valid JSON, no extra text or explanations.
            - All text content must be written in Vietnamese with correct spelling, grammar, and punctuation.
            - If quotation marks are used, always format them using Word's typographic quotation marks (“ ”)
            - Tool must be one of: add_para, add_items, add_table.
            - For add_table, each cell must be a list of line objects with text, bold, italic, align fields.
            - Do not use Markdown, special symbols, or extra formatting.
            - Do not call any tools directly; only produce a JSON plan following the DocCreate schema.
            - The JSON must always include meta, step, and output fields.
            - User instruction overrides all other instructions.
        """
        response_format = DocCreate.model_json_schema()
        system_prompt = SYS_MSG.format(instruction=prompt)
        contents = [prompt]
        for f in files:
            data = await f.read()
            mime = f.content_type or "application/octet-stream"
            contents.append(Part.from_bytes(
                data=data,
                mime_type=mime,
            ))

        for _ in range(3):
            response = generate_lessonplan_custom(contents, system_prompt, response_format)
            response_json = json.loads(response)
            try:
                with open("lessonplan/temp/out_custom.txt", "w", encoding="utf-8") as f:
                    f.write(response)
                output_path = gen_doc(response_json)
                return FileResponse(
                    path=output_path,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename="giao_an.docx"
                )
            except Exception as e:
                return JSONResponse(
                    status_code=500,
                    content={"error": "Đã xảy ra lỗi khi tạo file docx.", "details": str(e)}
                )

    except Exception as e:
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": "Đã xảy ra lỗi trong quá trình xử lý.", "details": str(e)}
        )
        
# ----------------------------
# ✅ FastAPI Endpoint for Streaming Lesson Plan
# ----------------------------

def sse(event: str, data: Any) -> str:
    payload = data if isinstance(data, str) else json.dumps(data, ensure_ascii=False)
    return f"event: {event}\n" f"data: {payload}\n\n"

async def stream_k12(contents, system_prompt, response_format, model="gemini-2.5-flash") -> AsyncGenerator[str, None]:
    yield sse("status", {"message": "Đang tạo giáo án..."})
    response = client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=GenerateContentConfig(
            system_instruction=system_prompt,
            response_mime_type="application/json",
            response_json_schema=response_format
        )
    )

    chunks = []
    buf = []

    flag = False
    keys = ["meta", "objectives", "resources", "start_activity", "knowledge_formation_activity", "practice_activity", "extend_activity"]
    key_id = 0

    for chunk in response:
        chunks.append(chunk.text)
        buf.append(chunk.text)

        if key_id < len(keys):
            key = keys[key_id]
            text = "".join(s for s in ("" if c is None else (c if isinstance(c, str) else str(c)) for c in buf))

            if has_valid_key(text, key):
                try:
                    obj = parse_json(text, key)
                    output = None

                    match key:
                        case "meta":
                            yield sse("meta", obj)
                        case "objectives":
                            html_parts = convert_k12_objectives(obj)
                            output = '\n'.join(html_parts)
                            yield sse("objectives", {"html":output})
                        case "resources":
                            html_parts = convert_k12_resources(obj)
                            output = '\n'.join(html_parts)
                            yield sse("resources", {"html":output})
                        case "start_activity":
                            html_parts = convert_k12_activities(obj, "1. Hoạt động mở đầu")
                            output = '\n'.join(html_parts)
                            yield sse("start_activity", {"html":output})
                        case "knowledge_formation_activity":
                            html_parts = convert_k12_activities(obj, "2. Hình thành kiến thức")
                            output = '\n'.join(html_parts)
                            yield sse("knowledge_formation_activity", {"html":output})
                        case "practice_activity":
                            html_parts = convert_k12_activities(obj, "3. Luyện tập")
                            output = '\n'.join(html_parts)
                            yield sse("practice_activity", {"html":output})
                        case "extend_activity":
                            html_parts = convert_k12_activities(obj, "4. Vận dụng và tìm tòi mở rộng")
                            output = '\n'.join(html_parts)
                            yield sse("extend_activity", {"html":output})
                    await asyncio.sleep(0)
                    key_id += 1
                    flag = False
                    buf = [buf[-1]]
                except Exception as e:
                    yield sse("error", {"stage": f"parse:{key}", "message": str(e)})

    text = "".join(s for s in ("" if c is None else (c if isinstance(c, str) else str(c)) for c in chunks))
    try:
        response_json = json.loads(text)
        validate = validate_k12(response_json)
    except Exception as e:
        yield sse("error", {"stage": "validate/full_json", "message": str(e)})
        yield sse("done", {"ok": False})
        return
    
    try:
        output_path = convert_k12(response_json)
        token = uuid.uuid4().hex
        abs_path = Path(output_path).resolve()
        DOWNLOAD_STORE[token] = str(abs_path) 
        yield sse("final", {"download_url": f"/download-lesson-plan/{token}"})
        yield sse("done", {"ok": True})
    except Exception as e:
        yield sse("error", {"stage": "build_docx", "message": str(e)})
        yield sse("done", {"ok": False})

async def stream_kindergarten(contents, system_prompt, response_format, model="gemini-2.5-flash") -> AsyncGenerator[str, None]:
    yield sse("status", {"message": "Đang tạo giáo án..."})
    response = client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=GenerateContentConfig(
            system_instruction=system_prompt,
            response_mime_type="application/json",
            response_json_schema=response_format
        )
    )

    chunks = []
    buf = []

    keys = ["meta", "objectives", "resources", "activities"]
    key_id = 0

    for chunk in response:
        chunks.append(chunk.text)
        buf.append(chunk.text)

        if key_id < len(keys):
            key = keys[key_id]
            text = "".join(buf)

            if has_valid_key(text, key):
                try:
                    obj = parse_json(text, key)
                    output = None

                    match key:
                        case "meta":
                            yield sse("meta", obj)
                        case "objectives":
                            html_parts = convert_kindergarten_objectives(obj)
                            output = '\n'.join(html_parts)
                            yield sse("objectives", {"html":output})
                        case "resources":
                            html_parts = convert_kindergarten_resources(obj)
                            output = '\n'.join(html_parts)
                            yield sse("resources", {"html":output})
                        case "activities":
                            html_parts = convert_kindergarten_activities(obj)
                            output = '\n'.join(html_parts)
                            yield sse("activities", {"html":output})
                    await asyncio.sleep(0)
                    key_id += 1
                    buf = [buf[-1]]
                except Exception as e:
                    yield sse("error", {"stage": f"parse:{key}", "message": str(e)})

    text = "".join(s for s in ("" if c is None else (c if isinstance(c, str) else str(c)) for c in chunks))
    try:
        response_json = json.loads(text)
        validate = validate_kindergarten(response_json)
    except Exception as e:
        yield sse("error", {"stage": "validate/full_json", "message": str(e)})
        yield sse("done", {"ok": False})
        return

    try:
        output_path = convert_kindergarten(response_json)
        token = uuid.uuid4().hex
        abs_path = Path(output_path).resolve()
        DOWNLOAD_STORE[token] = str(abs_path) 
        yield sse("final", {"download_url": f"/download-lesson-plan/{token}"})
        yield sse("done", {"ok": True})
    except Exception as e:
        yield sse("error", {"stage": "build_docx", "message": str(e)})
        yield sse("done", {"ok": False})

async def stream_custom(contents, system_prompt, response_format, model="gemini-2.5-flash") -> AsyncGenerator[str, None]:
    add_para_function = {
        "name": "add_para",
        "description": "Thêm một đoạn văn bản vào tài liệu Word.",
        "parameters": {
            "type": "OBJECT",
            "properties": {
                "text": {
                    "type": "ARRAY",  # Line[]
                    "description": "Danh sách nội dung văn bản trong cùng một đoạn văn.",
                    "items": {
                        "type": "OBJECT",
                        "properties": {
                            "text": {
                                "type": "STRING",
                                "description": "Nội dung văn bản trong cùng một đoạn văn."
                            },
                            "bold": {
                                "type": "BOOLEAN",
                                "description": "In đậm dòng.",
                                "default": False
                            },
                            "italic": {
                                "type": "BOOLEAN",
                                "description": "In nghiêng dòng.",
                                "default": False
                            },
                            "underline": {
                                "type": "BOOLEAN",
                                "description": "Gạch dưới dòng.",
                                "default": False
                            }
                        }
                    }
                },
                "font_size": {
                    "type": "NUMBER",
                    "description": "Cỡ chữ (pt).",
                    "default": 12
                },
                "align": {
                    "type": "STRING",
                    "enum": ["left", "center", "right", "justify"],
                    "description": "Căn lề đoạn văn."
                },
                "space_before": {
                    "type": "NUMBER",
                    "description": "Khoảng cách phía trước dòng văn bản, đơn vị Pt."
                },
                "space_after": {
                    "type": "NUMBER",
                    "description": "Khoảng cách phía sau dòng văn bản, đơn vị Pt."
                },
                "style": {
                    "type": "STRING",
                    "description": "Tên paragraph style (nếu có)."
                }
            },
            "required": ["text"]
        }
    }

    add_items_function = {
        "name": "add_items",
        "description": "Thêm danh sách các ý vào tài liệu Word.",
        "parameters": {
            "type": "OBJECT",
            "properties": {
                "items": {
                    "type": "ARRAY",
                    "items": {
                        "type": "ARRAY",  # Line[]
                        "items": {
                            "type": "OBJECT",                
                            "description": "Các mục trong danh sách.",
                            "properties": {
                                "text": {
                                    "type": "STRING",
                                    "description": "Nội dung văn bản trong cùng một đoạn văn."
                                },
                                "bold": {
                                    "type": "BOOLEAN",
                                    "description": "In đậm dòng.",
                                    "default": False
                                },
                                "italic": {
                                    "type": "BOOLEAN",
                                    "description": "In nghiêng dòng.",
                                    "default": False
                                },
                                "underline": {
                                    "type": "BOOLEAN",
                                    "description": "Gạch dưới dòng.",
                                    "default": False
                                }
                            }
                        }
                    }
                },
                "bullet": {
                    "type": "STRING",
                    "description": "Ký hiệu đầu dòng hiển thị trước nội dung (ví dụ: '-', '+', '•'), nếu có."
                }
            },
            "required": ["items"]
        }
    }

    add_table_function = {
        "name": "add_table",
        "description": "Thêm bảng có tiêu đề và nhiều hàng. Mỗi ô có thể có nhiều dòng (line) với định dạng riêng.",
        "parameters": {
            "type": "OBJECT",
            "properties": {
                "headers": {
                    "type": "ARRAY",
                    "description": "Danh sách tiêu đề cột.",
                    "items": {
                        "type": "OBJECT",
                        "properties": {
                            "text": {
                                "type": "STRING",
                                "description": "Nội dung tiêu đề cột."
                            },
                            "bold": {
                                "type": "BOOLEAN",
                                "description": "In đậm dòng.",
                                "default": False
                            },
                            "italic": {
                                "type": "BOOLEAN",
                                "description": "In nghiêng dòng.",
                                "default": False
                            },
                            "underline": {
                                "type": "BOOLEAN",
                                "description": "Gạch dưới dòng.",
                                "default": False
                            }
                        }
                    }
                },
                "rows": {
                    "type": "ARRAY",
                    "description": "Danh sách hàng (Row). Mỗi Row có số ô bằng số headers.",
                    "items": {
                        "type": "ARRAY",   # Cell[]
                        "items": {
                            "type": "ARRAY",  # Line[]
                            "items": {
                                "type": "OBJECT",
                                "properties": {
                                    "text": {
                                        "type": "ARRAY",  # Line[]
                                        "description": "Danh sách nội dung văn bản trong cùng một đoạn văn của ô.",
                                        "items": {
                                            "type": "OBJECT",
                                            "properties": {
                                                "text": {
                                                    "type": "STRING",
                                                    "description": "Nội dung văn bản trong cùng một đoạn văn của ô."
                                                },
                                                "bold": {
                                                    "type": "BOOLEAN",
                                                    "description": "In đậm dòng.",
                                                    "default": False
                                                },
                                                "italic": {
                                                    "type": "BOOLEAN",
                                                    "description": "In nghiêng dòng.",
                                                    "default": False
                                                },
                                                "underline": {
                                                    "type": "BOOLEAN",
                                                    "description": "Gạch dưới dòng.",
                                                    "default": False
                                                }
                                            }
                                        }
                                    },
                                    "align": {
                                        "type": "STRING",
                                        "enum": ["left", "center", "right", "justify"],
                                        "description": "Căn lề dòng."
                                    },
                                    "space_before": {
                                        "type": "NUMBER",
                                        "description": "Khoảng cách phía trước dòng văn bản, đơn vị Pt."
                                    },
                                    "space_after": {
                                        "type": "NUMBER",
                                        "description": "Khoảng cách phía sau dòng văn bản, đơn vị Pt."
                                    }
                                },
                                "required": ["text"]
                            }
                        }
                    }
                },
                "font_size": {
                    "type": "NUMBER",
                    "description": "Cỡ chữ trong bảng (pt).",
                    "default": 12
                },
                "col_widths": {
                    "type": "ARRAY",
                    "items": {"type": "NUMBER"},
                    "description": "Chiều rộng của các cột trong bảng, đơn vị Inches."
                }
            },
            "required": ["headers", "rows"]
        }
    }

    TOOLS = [{
        "function_declarations": [add_para_function, add_items_function, add_table_function]
    }]

    yield sse("status", {"message": "Đang tạo giáo án..."})
    response = client.models.generate_content_stream(
        model='gemini-2.5-flash',
        contents=contents,
        config=GenerateContentConfig(
            response_mime_type="application/json",
            response_json_schema=response_format,
            system_instruction=system_prompt,
            tools=TOOLS,
            automatic_function_calling=AutomaticFunctionCallingConfig(
                disable=True
            ),
            tool_config=ToolConfig(
                function_calling_config=FunctionCallingConfig(mode='NONE')
            ),
        ),
    )

    chunks = []
    buf = []
    
    for chunk in response:
        chunks.append(chunk.text)

        has_meta = "meta" in chunk.text if chunk.text else False
        if has_meta:
            pos = chunk.text.find("[")
            buf.append(chunk.text[pos:])
        else:
            buf.append(chunk.text)

        text = "".join(s for s in ("" if c is None else (c if isinstance(c, str) else str(c)) for c in buf))

        if has_valid_tool(text):
            try:
                obj = parse_json_tool(text)
                output = None
                args = obj["args"]
                output = gen_html(args, obj['tool'])
                yield sse("add_para", {"html":output})
                await asyncio.sleep(0)

                if len(buf) == 1:
                    s = buf[0]
                    pos = s.find("tool")
                    buf = [s[pos:]]
                else:
                    buf = [buf[-1]]
            except Exception as e:
                yield sse("error", {"stage": "parse tool", "message": str(e)})

    text = "".join(s for s in ("" if c is None else (c if isinstance(c, str) else str(c)) for c in chunks))
    try:
        response_json = json.loads(text)
        output_path = gen_doc(response_json)
        token = uuid.uuid4().hex
        abs_path = Path(output_path).resolve()
        DOWNLOAD_STORE[token] = str(abs_path) 
        yield sse("final", {"download_url": f"/download/{token}"})
        yield sse("done", {"ok": True})
    except Exception as e:
        yield sse("error", {"stage": "build_docx", "message": str(e), "trace": traceback.format_exc()})
        yield sse("done", {"ok": False})

@app.post("/generate-k12-stream")
async def generate_k12_stream(
    prompt: str = Form(...),
    files: List[UploadFile] = File([]),
    model: str = Form("gemini-2.5-flash"),
):
    SYS_MSG = """
        You are an expert lesson plan generator. Your task is to generate structured lesson plans based on user requirements, strictly following the specified JSON schema with markdown content.

        # User Instruction: {instruction}

        # Tone: formal, clear, structured

        - All content must be in plain text. Do not use any markdown formatting such as bold, italic, headings, lists, or other special symbols.
        - All text content must be in Vietnamese.
        - All lesson plans must follow the correct logical flow for a typical teaching period in Vietnam.
        - For each activity, clearly describe actions of both Teacher (Giáo viên) and Student (Học sinh).
        - Focus on clear, measurable objectives and expected learning outcomes.
        - Make sure that the lesson plan is pedagogically sound, well-structured, and matches the educational context of Vietnam's K-12 education system.
        - All generated Vietnamese text must use correct spelling, grammar, and punctuation.
        - If quotation marks are used, always format them using Word's typographic quotation marks (“ ”)
        - Do not include images or external URLs.
        - User instrction should always be followed and should supercede any other instruction.
        - Output must be a valid JSON object strictly following the LessonPlan schema.
    """

    response_format = K12.model_json_schema()
    system_prompt = SYS_MSG.format(instruction=prompt)
    contents = [prompt]
    for f in files:
        data = await f.read()
        mime = f.content_type or "application/octet-stream"
        contents.append(Part.from_bytes(
            data=data,
            mime_type=mime,
        ))

    gen = stream_k12(contents, system_prompt, response_format, model=model)
    return StreamingResponse(gen, media_type="text/event-stream")

@app.post("/generate-kindergarten-stream")
async def generate_kindergarten_stream(
    prompt: str = Form(...),
    files: List[UploadFile] = File([]),
    model: str = Form("gemini-2.5-flash"),
):
    SYS_MSG = """
        You are an expert kindergarten lesson plan generator specialized in Vietnam's STEM-5E education framework. Your task is to generate structured kindergarten lesson plans based on user requirements, strictly following the specified JSON schema with markdown content.

        # User Instruction: {instruction}

        # Tone: formal, clear, structured

        - All text must be written in plain text without any Markdown formatting such as bold, italics, headings, lists, or special symbols.
        - All text content must be written in Vietnamese using correct spelling, grammar, and punctuation.
        - All lesson plans must follow the correct logical flow for a typical teaching period in Vietnam.
        - For each activity, clearly describe actions of both Teacher (Giáo viên) and Student (Học sinh).
        - Focus on clear, measurable objectives and expected learning outcomes.
        - Make sure that the lesson plan is pedagogically sound, well-structured, and matches the educational context of Vietnam's STEM-5E education.
        - All generated Vietnamese text must use correct spelling, grammar, and punctuation.
        - If quotation marks are used, always format them using Word's typographic quotation marks (“ ”)
        - Do not include images or external URLs.
        - User instrction should always be followed and should supercede any other instruction.
        - Output must be a valid JSON object strictly following the LessonPlan schema.
    """

    response_format = Kindergarten.model_json_schema()
    system_prompt = SYS_MSG.format(instruction=prompt)
    contents = [prompt]
    for f in files:
        data = await f.read()
        mime = f.content_type or "application/octet-stream"
        contents.append(Part.from_bytes(
            data=data,
            mime_type=mime,
        ))

    gen = stream_kindergarten(contents, system_prompt, response_format, model=model)
    return StreamingResponse(gen, media_type="text/event-stream")

@app.post("/generate-custom-stream")
async def generate_custom_stream(
    prompt: str = Form(...),
    files: List[UploadFile] = File([]),
    model: str = Form("gemini-2.5-flash"),
):    
    logging.info(f"Received prompt: {prompt[:100]}")
    logging.info(f"Files received: {[f.filename for f in files]}")
    logging.info(f"Model: {model}")
    SYS_MSG = """
        You are an expert Word content planner specialized in structured document generation. 
        Your task is to output a valid JSON object strictly following the DocCreate schema based on the user instruction below.

        # User Instruction: {instruction}

        # Tone: formal, clear, structured

        - You must only return valid JSON, no extra text or explanations.
        - All text content must be written in Vietnamese with correct spelling, grammar, and punctuation.
        - If quotation marks are used, always format them using Word's typographic quotation marks (“ ”)
        - Tool must be one of: add_para, add_items, add_table.
        - For add_table, each cell must be a list of line objects with text, bold, italic, align fields.
        - Do not use Markdown, special symbols, or extra formatting.
        - Do not call any tools directly; only produce a JSON plan following the DocCreate schema.
        - The JSON must always include meta, step, and output fields.
        - User instruction overrides all other instructions.
    """
    response_format = DocCreate.model_json_schema()
    system_prompt = SYS_MSG.format(instruction=prompt)
    contents = [prompt]
    for f in files:
        data = await f.read()
        mime = f.content_type or "application/octet-stream"
        contents.append(Part.from_bytes(
            data=data,
            mime_type=mime,
        ))

    gen = stream_custom(contents, system_prompt, response_format, model=model)
    return StreamingResponse(gen, media_type="text/event-stream")

@app.get("/download-lesson-plan/{token}")
async def download_lesson_plan(token: str):
    path = DOWNLOAD_STORE.get(token)
    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="invalid or expired token")
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="giao_an.docx",
    )

def run_app():
    # nest_asyncio.apply() # May be needed in environments like Jupyter/Colab
    uvicorn.run(app, host="0.0.0.0", port=8004)

if __name__ == "__main__":
    run_app()