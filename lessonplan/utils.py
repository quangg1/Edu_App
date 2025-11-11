import json, re, os

from docxtpl import DocxTemplate, RichText
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .model.k12_model import K12
from .model.kindergarten_model import Kindergarten

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "template")
RESULT_DIR = os.path.join(BASE_DIR, "result")

def set_font_size(container, font_size=12):
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)

def clear_paragraph_format(p):
    """Xóa thụt đầu dòng và khoảng cách trong paragraph"""
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def reset_document_format(doc):
    """
    Reset toàn bộ định dạng mặc định sau 1 vòng lặp:
    - Xóa numbering
    - Xóa indent
    - Xóa khoảng cách
    """
    for p in doc.paragraphs:
        clear_paragraph_format(p)

def set_cell_borders(cell):
    """Thêm đường viền cho một ô trong bảng"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')   
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') 
        tcBorders.append(border)

    tcPr.append(tcBorders)

def add_lines_to_cell(cell, lines, bold_first_line=False, font_name="Times New Roman", font_size=12):
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0]
        clear_paragraph_format(p)
        run = p.add_run(line+"\n")
        if bold_first_line and idx == 0:
            run.bold = True
        run.font.name = font_name
        run.font.size = Pt(font_size)

    if p.runs and p.runs[-1].text.endswith("\n"):
        p.runs[-1].text = p.runs[-1].text.rstrip("\n")

def create_activity_table(container, step_data, order, cols, hdr_text, font_size=12, font_name="Times New Roman"):
    table = container.add_table(rows=1, cols=cols)
    table.autofit = True

    hdr_cells = table.rows[0].cells

    for i, text in enumerate(hdr_text):
        p = hdr_cells[i].paragraphs[0]
        clear_paragraph_format(p)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = font_name
        set_cell_borders(hdr_cells[i])

    for key, title in order:
        block = step_data.get(key)
        if not block:
            continue

        left_text = []
        right_text = []

        for field_name, field_value in block.items():
            if field_name in ("activity", "teacher_activity"):  
                if isinstance(field_value, list):
                    left_text.extend(field_value)
                else:
                    left_text.append(str(field_value))
            elif field_name in ("expected_products", "student_activity"):
                if isinstance(field_value, list):
                    right_text.extend(field_value)
                else:
                    right_text.append(str(field_value))

        left_lines = [title] + left_text
        right_lines = right_text

        row_cells = table.add_row().cells
        add_lines_to_cell(row_cells[0], left_lines, bold_first_line=True,
                          font_name=font_name, font_size=font_size)
        add_lines_to_cell(row_cells[1], right_lines, bold_first_line=False,
                          font_name=font_name, font_size=font_size)
        
        for c in row_cells:
            set_cell_borders(c)

    return table

def main_activity_to_text(tpl, data):
    """
    tpl: DocxTemplate object
    data: dict chứa nội dung phần Hình thành kiến thức (sub_activities)
    """
    sub = tpl.new_subdoc()

    sub_activities = data.get("sub_activities", [])

    for id, sub_act in enumerate(sub_activities):
        p = sub.add_paragraph()
        run = p.add_run(sub_act.get("title", ""))
        run.bold = True
        run.add_break()   

        gi = sub_act.get("general_info", {})
        if gi:
            run = p.add_run(f"Mục tiêu: ")
            run.bold = True
            p.add_run(f"{gi.get('objective','')}\n")
            run = p.add_run(f"Nội dung: ")
            run.bold = True
            p.add_run(f"{gi.get('content','')}\n")
            run = p.add_run(f"Sản phẩm: ")
            run.bold = True
            p.add_run(f"{gi.get('expected_products','')}\n")
            run = p.add_run(f"Tổ chức thực hiện: ")
            run.bold = True

        step_data = sub_act.get("step", {})

        if step_data:
            order = [
                ("assign_task",   "Bước 1: Chuyển giao nhiệm vụ"),
                ("do_task",       "Bước 2: Thực hiện nhiệm vụ"),
                ("discuss_task",  "Bước 3: Báo cáo, thảo luận"),
                ("conclusion_task","Bước 4: Kết luận, nhận định"),
            ]
            hdr_text = ["Hoạt động của GV và HS", "Sản phẩm dự kiến"]
            create_activity_table(sub, step_data, order, 2, hdr_text, font_size=12, font_name="Times New Roman")

        if id < len(sub_activities) - 1:
            blank = sub.add_paragraph()

        reset_document_format(sub)

    set_font_size(sub, 12)
    
    return sub

def activity_to_text(activities):
    act = RichText()

    act.add("Mục tiêu:", bold=True, size=24)
    act.add(f" {activities['general_info']['objective']}\n", size=24)

    act.add("Nội dung:", bold=True, size=24)
    act.add(f" {activities['general_info']['content']}\n", size=24)

    act.add("Sản phẩm:", bold=True, size=24)
    act.add(f" {activities['general_info']['expected_products']}\n", size=24)

    act.add("Tổ chức thực hiện\n", bold=True, size=24)
    for step in activities['execution']['execute_description']:
        act.add(f"– {step}\n", size=24)

    execution = activities.get('execution', {})
    teacher_prompts = execution.get('teacher_prompts', [])
    if teacher_prompts:
        act.add("Câu hỏi gợi ý\n", bold=True, size=24)
        for step in teacher_prompts:
            act.add(f"– {step}\n", size=24)
            
    return act

def list_to_text(lst):
    return "\n".join([f"– {item}" for item in lst]) if lst else ""

def convert_k12(data, template="k12_template.docx", output="giao_an_k12.docx"):
    template_path = os.path.join(TEMPLATE_DIR, template)
    tpl = DocxTemplate(template_path)

    meta = data.get("meta", {})
    objectives = data.get("objectives", {})
    resources = data.get("resources", {})
    start_activity = data.get("start_activity", {})
    knowledge_formation_activity = data.get("knowledge_formation_activity", {})
    practice_activity = data.get("practice_activity", {})
    extend_activity = data.get("extend_activity", {})

    knowledge_text = list_to_text(objectives.get("knowledge"))
    competencies_text = list_to_text(objectives.get("competencies"))
    qualities_text = list_to_text(objectives.get("qualities"))

    teacher_materials = resources.get("teacher", [])
    student_materials = resources.get("student", [])

    resource_teacher_text =  list_to_text(teacher_materials)
    resource_student_text =  list_to_text(student_materials)

    mo_dau = activity_to_text(start_activity)
    hinh_thanh_kien_thuc = main_activity_to_text(tpl, knowledge_formation_activity)
    luyen_tap = activity_to_text(practice_activity)
    van_dung_mo_rong = activity_to_text(extend_activity)

    context = {
        "ten_truong": (meta.get("school_name") or "").upper(),
        "gv": (meta.get("teacher_name") or "").upper(),
        "to_cm": (meta.get("department") or "").upper(),
        "ten_bai_day": meta.get("lesson_title", ""),
        "mon_hoc": meta.get("subject", ""),
        "lop": meta.get("grade", ""),
        "so_tiet": meta.get("periods", ""),

        "muc_tieu_kien_thuc": knowledge_text,
        "muc_tieu_nang_luc": competencies_text,
        "muc_tieu_pham_chat": qualities_text,
        "resource_gv": resource_teacher_text,
        "resource_hs": resource_student_text,

        "mo_dau": mo_dau,
        "hinh_thanh_kien_thuc": hinh_thanh_kien_thuc,
        "luyen_tap": luyen_tap,
        "van_dung_mo_rong": van_dung_mo_rong
    }

    output_path = os.path.join(RESULT_DIR, output)

    tpl.render(context)
    tpl.save(output_path)

    return output_path

def convert_kindergarten(data, template="kindergarten_template.docx", output="giao_an_mam_non.docx"):
    template_path = os.path.join(TEMPLATE_DIR, template)
    tpl = DocxTemplate(template_path)

    meta = data.get("meta", {})
    objectives = data.get("objectives", {})
    resources = data.get("resources", {})
    activities = data.get("activities", {})

    teacher_materials = resources.get("teacher", [])
    student_materials = resources.get("student", [])

    resource_teacher_text =  list_to_text(teacher_materials)
    resource_student_text =  list_to_text(student_materials)

    if activities:
        to_chuc_hoat_dong = tpl.new_subdoc()
        order = [
            ("engagement",   "Bước 1: Gắn kết"),
            ("exploration",       "Bước 2: Khám phá"),
            ("explanation",  "Bước 3: Giải thích và chia sẻ"),
            ("elaborate","Bước 4: Áp dụng"),
            ("evaluate","Bước 5: Đánh giá"),
        ]
        hdr_text = ["Hoạt động của giáo viên", "Hoạt động của trẻ"]
        create_activity_table(to_chuc_hoat_dong, activities, order, 2, hdr_text, font_size=12, font_name="Times New Roman")

    context = {
        "ten_truong": (meta.get("school_name") or "").upper(),
        "gv": (meta.get("teacher_name") or ""),
        "chu_de": (meta.get("subject", "") or "").upper(),
        "de_tai": meta.get("lesson_title", ""),

        "khoa_hoc": list_to_text(objectives.get("steam", []).get("science", [])),
        "cong_nghe": list_to_text(objectives.get("steam", []).get("technology", [])),
        "ky_thuat": list_to_text(objectives.get("steam", []).get("engineering", [])),
        "toan_hoc": list_to_text(objectives.get("steam", []).get("mathematics", [])),
        "muc_tieu_ky_nang": list_to_text(objectives.get("skills")),
        "muc_tieu_pham_chat": list_to_text(objectives.get("attitude")),

        "dia_diem": (resources.get("place") or ""),
        "resource_gv": resource_teacher_text,
        "resource_hs": resource_student_text,

        "to_chuc_hoat_dong": to_chuc_hoat_dong
    }

    output_path = os.path.join(RESULT_DIR, output)

    tpl.render(context)
    tpl.save(output_path)

    return output_path

def validate_k12(data: dict) -> K12:
    return K12.model_validate(data)

def validate_kindergarten(data: dict) -> Kindergarten:
    return Kindergarten.model_validate(data)

def has_valid_key(text: str, key: str):
    pattern = fr'"{re.escape(key)}"\s*:\s*[\{{\[]'
    m = re.search(pattern, text)
    if not m:
        return False
    
    start_pos = text.find(m.group()[-1], m.end() - 1)
    
    decoder = json.JSONDecoder()

    try:
        obj, end_pos = decoder.raw_decode(text, start_pos)
        return True
    except json.JSONDecodeError:
        return False
    
def has_valid_tool(text: str):
    pattern = r'\{\s*"tool":'
    m = re.search(pattern, text)
    if not m:
        return False
    
    start_pos = m.start()

    decoder = json.JSONDecoder()

    try:
        obj, end_pos = decoder.raw_decode(text, start_pos)
        return True
    except json.JSONDecodeError:
        return False
    
def parse_json(text: str, key: str):
    pattern = fr'"{re.escape(key)}"\s*:\s*[\{{\[]'
    m = re.search(pattern, text)
    
    start_pos = text.find(m.group()[-1], m.end() - 1)

    decoder = json.JSONDecoder()

    obj, end_pos = decoder.raw_decode(text, start_pos)

    return obj

def parse_json_tool(text: str):
    pattern = r'\{\s*"tool":'
    m = re.search(pattern, text)
    
    start_pos = m.start()

    decoder = json.JSONDecoder()

    obj, end_pos = decoder.raw_decode(text, start_pos)
    
    return obj