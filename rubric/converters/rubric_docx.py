# converters/rubric_docx.py
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ─────────────────────────────────────────────────────────────────────────────
# Helpers (small utilities)
# ─────────────────────────────────────────────────────────────────────────────

def _set_heading(p, text: str, size: int = 16, bold: bool = True) -> None:
    """Write a heading-like paragraph."""
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)

def _cell_set(p, text: str, *, bold: bool = False, size: int = 10, center: bool = False) -> None:
    """Write text into a table cell paragraph with simple formatting."""
    # Clear existing runs to avoid mixed styling
    for r in list(p.runs):
        r.clear()
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _infer_levels_from_first_criterion(criteria: List[Dict]) -> List[str]:
    """
    If scale.levels is missing/empty, infer level labels from the first criterion's levels.
    Fallback to common Vietnamese labels when nothing is available.
    """
    if not criteria:
        return ["Xuất sắc", "Tốt", "Đạt", "Cần cải thiện"]
    first = criteria[0] or {}
    levels = first.get("levels") or []
    labels = [lv.get("label") for lv in levels if lv and lv.get("label")]
    return labels or ["Xuất sắc", "Tốt", "Đạt", "Cần cải thiện"]

# ─────────────────────────────────────────────────────────────────────────────
# Core converter
# ─────────────────────────────────────────────────────────────────────────────

def rubric_to_docx(
    rubric: dict,
    out_path: str,
    *,
    landscape: bool = True,
    margins_inch: Dict[str, float] = None
) -> str:
    """
    Convert a rubric JSON (matching your RubricSchema) into a .docx file.

    Layout:
      - Title line: "Bảng Rubrics đánh giá: <rubric_title>"
      - Meta line: "Môn học | Khối lớp | Hình thức"
      - Table:
          Row 0: merged title "Bảng Rubrics đánh giá"
          Row 1: headers: "Mức độ" + each criterion "Name (weight%)"
          Rows 2..: each level (scale.levels) per row, each cell is a description matched by level label.

    Args:
        rubric: Rubric JSON as python dict (must follow your RubricSchema).
        out_path: Path to save the .docx output.
        landscape: Whether to set page orientation to landscape (default: True).
        margins_inch: Optional custom margins (dict with keys left/right/top/bottom in inches).

    Returns:
        The same out_path (string).
    """
    title       = rubric.get("rubric_title", "Rubric")
    subject     = rubric.get("subject", "")
    grade_level = rubric.get("grade_level", "")
    assess_type = rubric.get("assessment_type", "") or ""
    criteria    = rubric.get("criteria", []) or []
    scale       = rubric.get("scale", {}) or {}
    level_names = scale.get("levels") or _infer_levels_from_first_criterion(criteria)

    # Build grid of descriptions per level per criterion
    # grid[row(level_idx)][col(criterion_idx)] = description
    grid: List[List[str]] = []
    for lv in level_names:
        row_cells: List[str] = []
        for crit in criteria:
            desc = ""
            for lv_obj in (crit.get("levels") or []):
                if lv_obj.get("label") == lv:
                    desc = lv_obj.get("description", "")
                    break
            row_cells.append(desc)
        grid.append(row_cells)

    # Create document
    doc = Document()

    # Default font to render Vietnamese properly
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    # Ensure East Asian font mapping to avoid wrong glyphs
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    # Page settings (orientation + margins)
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap width/height to apply landscape
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    # Reasonable default margins for print-friendly layout
    if margins_inch is None:
        margins_inch = {"left": 0.8, "right": 0.8, "top": 0.6, "bottom": 0.6}
    section.left_margin = Inches(margins_inch.get("left", 0.8))
    section.right_margin = Inches(margins_inch.get("right", 0.8))
    section.top_margin = Inches(margins_inch.get("top", 0.6))
    section.bottom_margin = Inches(margins_inch.get("bottom", 0.6))

    # Title
    p_title = doc.add_paragraph()
    _set_heading(p_title, f"Bảng Rubrics đánh giá: {title}", size=16, bold=True)

    # Meta
    meta = f"Môn học: {subject}    Khối lớp: {grade_level}    Hình thức: {assess_type}"
    doc.add_paragraph(meta)
    doc.add_paragraph("")  # spacer

    # Table structure
    cols = 1 + len(criteria)      # "Mức độ" + each criterion
    rows = 2 + len(level_names)   # 2 header rows + N level rows
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    # Header row 0 - merged title
    hdr0 = table.rows[0].cells
    if len(hdr0) > 1:
        hdr0[0].merge(hdr0[-1])
    _cell_set(hdr0[0].paragraphs[0], "Bảng Rubrics đánh giá", bold=True, size=12, center=True)

    # Header row 1 - column headers
    hdr1 = table.rows[1].cells
    _cell_set(hdr1[0].paragraphs[0], "Mức độ", bold=True, center=True)
    for j, crit in enumerate(criteria, start=1):
        name = crit.get("name", "Tiêu chí")
        w    = crit.get("weight_percent", 0)
        _cell_set(hdr1[j].paragraphs[0], f"{name} ({int(w)}%)", bold=True, center=True)

    # Body rows - per level
    for i, lv in enumerate(level_names, start=2):
        row_cells = table.rows[i].cells
        _cell_set(row_cells[0].paragraphs[0], lv, bold=True, center=True)
        for j, desc in enumerate(grid[i-2], start=1):
            _cell_set(row_cells[j].paragraphs[0], desc, size=10, bold=False)

    # Gentle spacing for readability
    for r in table.rows:
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(4)

    # (Optional) Let Word auto-fit content; python-docx doesn't expose direct column width control reliably.
    try:
        table.autofit = True
    except Exception:
        pass

    doc.save(out_path)
    return out_path
